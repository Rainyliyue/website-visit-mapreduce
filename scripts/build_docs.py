from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, after=0, line_spacing=1.05)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "D9E2EC")
        borders.append(border)
    tbl_pr.append(borders)


def apply_document_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("MapReduce 课程设计")
    set_run_font(run, size=9, color="666666")


def add_code_block(document, lines):
    for line in lines:
        p = document.add_paragraph()
        set_paragraph_format(p, after=2, line_spacing=1.0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)


def add_table(document, table_lines):
    rows = []
    for line in table_lines:
        stripped = line.strip().strip("|")
        rows.append([cell.strip() for cell in stripped.split("|")])
    if len(rows) < 2:
        return
    separator = rows[1]
    if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in separator):
        return
    data_rows = [rows[0], *rows[2:]]
    table = document.add_table(rows=len(data_rows), cols=len(data_rows[0]))
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(data_rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "F2F4F7")
    document.add_paragraph()


def parse_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        set_run_font(run, size=11)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def markdown_to_docx(md_path, docx_path):
    document = Document()
    apply_document_styles(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(paragraph, before=0, after=10, line_spacing=1.1)
            run = paragraph.add_run(stripped[2:])
            set_run_font(run, size=20, bold=True, color="0B2545")
        elif stripped.startswith("## "):
            paragraph = document.add_paragraph(stripped[3:], style="Heading 1")
            set_paragraph_format(paragraph, before=16, after=8, line_spacing=1.1)
        elif stripped.startswith("### "):
            paragraph = document.add_paragraph(stripped[4:], style="Heading 2")
            set_paragraph_format(paragraph, before=12, after=6, line_spacing=1.1)
        elif stripped.startswith("#### "):
            paragraph = document.add_paragraph(stripped[5:], style="Heading 3")
            set_paragraph_format(paragraph, before=8, after=4, line_spacing=1.1)
        elif re.match(r"^[-*]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            set_paragraph_format(paragraph, after=4, line_spacing=1.15)
            parse_inline_runs(paragraph, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            set_paragraph_format(paragraph, after=4, line_spacing=1.15)
            parse_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            paragraph = document.add_paragraph()
            set_paragraph_format(paragraph, after=6, line_spacing=1.1)
            parse_inline_runs(paragraph, stripped)
        i += 1

    document.save(docx_path)


def main():
    markdown_to_docx(DOCS / "需求分析与系统设计说明书.md", DOCS / "需求分析与系统设计说明书.docx")
    markdown_to_docx(DOCS / "课程设计报告.md", DOCS / "课程设计报告.docx")


if __name__ == "__main__":
    sys.exit(main())
